from flask import Flask, render_template, request, send_file, jsonify
import os
from docx import Document
import re

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Keyword validation
KEYWORDS = {
    "rise_tool": ["Lesson 1 of"],
    "storyline_tool": ["Translation"],
    "content_builder_tool": ["Screen Title", "Button"],
    "html_tool": ["f_preload_data"]
}

# Tool functions (implement your logic here)
def rise_tool(input_path, output_path):
    # Dummy processing logic
        def copy_and_filter_text_with_section_removals(input_path, output_path):
            input_doc = Document(input_path)
            output_doc = Document()
            regex_to_ignore = re.compile(r"Lesson \d+ of \d+|\b\d+ of \d+\b", re.IGNORECASE)
            exclude_keywords = ["CONT I NU E"]
            sections_to_remove = ["Navigation" ]
            in_section_to_remove = False
            title_captured = False
            start_copying = False

            for paragraph in input_doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                if not title_captured and any(run.bold and run.font.size and run.font.size.pt > 18 for run in paragraph.runs):
                    title_captured = True
                    new_paragraph = output_doc.add_paragraph()
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.bold = run.bold  # Retain bold styling
                        if run.font.size and run.font.size.pt > 18:
                            new_run.bold = True  # Ensure text is bold
                            new_run.underline = True  # Add underline to large text
                    continue  # Skip to the next paragraph after capturing the title

                if "lesson 1" in text.lower():
                    start_copying = True  # Start copying from this point onward

                if not start_copying:
                    continue
                
                if regex_to_ignore.search(text) or any(keyword in text for keyword in exclude_keywords):
                    continue  # Skip this paragraph

                if any(run.bold and run.font.size and run.font.size.pt > 18 for run in paragraph.runs):
                    if any(section.lower() in text.lower() for section in sections_to_remove):
                        in_section_to_remove = True
                        continue  # Skip the section heading itself
                
                if in_section_to_remove:
                    # Check for the end of the section (e.g., next heading or empty paragraph)
                    if any(run.bold and run.font.size and run.font.size.pt > 18 for run in paragraph.runs):
                        in_section_to_remove = False
                    else:
                        continue
                
                new_paragraph = output_doc.add_paragraph()
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    new_run.bold = run.bold  # Retain bold styling
                    if run.font.size and run.font.size.pt > 18:
                        new_run.bold = True  # Ensure text is bold
                        new_run.underline = True  # Add underline to large text

            output_doc.save(output_path)
            print(f"Filtered and formatted text has been saved to {output_path}")

        input_file = input_path
        output_file = output_path

        copy_and_filter_text_with_section_removals(input_file, output_file)

        # pass
        # doc = Document(input_path)
        # doc.add_paragraph("Processed with Rise Tool")
        # doc.save(output_path)

def storyline_tool(input_path, output_path):
    # Dummy processing logic
        def extract_translation_content(input_file, output_file):
            doc = Document(input_file)
            new_doc = Document()
            first_translation_table_skipped = False
            text_box_pattern = re.compile(r"text box(\s*\d*)", re.IGNORECASE)
            sections_to_remove = []
            previous_translation_content = None
            print(sections_to_remove)
            for table in doc.tables:
                header_cells = table.rows[0].cells
                header_texts = [cell.text.lower() for cell in header_cells]

                if "translation" in header_texts and "type" in header_texts:
                    if not first_translation_table_skipped:
                        first_translation_table_skipped = True
                        continue

                    translation_index = header_texts.index("translation")
                    type_index = header_texts.index("type")

                    if any(any(phrase in row.cells[translation_index].text.lower() for phrase in sections_to_remove) for row in table.rows[1:]):
                        continue  # Skip this table if it contains any skip phrases in "Type"
                    
                    combined_text_box_content = []

                    for row in table.rows[1:]:  # Skip the header row
                        translation_content = row.cells[translation_index].text.strip()
                        type_content = row.cells[type_index].text.strip().lower()

                        if "alt text" in type_content.lower() or "alttext" in type_content.lower():
                            # print("alt text came. Its ignored")
                            continue  # Ignore this row if "Alt Text" is in "Type"
                        if "Hover" in type_content.lower():
                            # print("alt text came. Its ignored")
                            continue  # Ignore this row if "Alt Text" is in "Type"

                        if translation_content not in ["Next >", "< Back"] and translation_content != previous_translation_content:
                            if text_box_pattern.match(type_content):
                                combined_text_box_content.append(translation_content)
                            else:
                                para = new_doc.add_paragraph()
                                if "slide name" in type_content:
                                    run = para.add_run(translation_content)
                                    run.bold = True
                                else:
                                    para.add_run(translation_content)

                            previous_translation_content = translation_content

                    if combined_text_box_content:
                        single_line_content = " ".join(combined_text_box_content)
                        new_doc.add_paragraph(single_line_content)

            new_doc.save(output_file)
            print(f'Translation content extracted and saved to {output_file}')

        input_file = input_path
        output_file = output_path
        extract_translation_content(input_file, output_file)

        # pass
        # doc = Document(input_path)
        # doc.add_paragraph("Processed with Storyline Tool")
        # doc.save(output_path)

def content_builder_tool(input_path, output_path):
    # Dummy processing logic
        def extract_content_after_colon(input_file, output_file):
            # Load the Word document
            doc = Document(input_file)
            
            # Create a new Word document for the output
            new_doc = Document()
            content_added = False
            
            # Regular expression to detect HTML tags
            html_tag_pattern = re.compile(r"<[^>]+>")
            
            # Loop through each paragraph in the document
            for para in doc.paragraphs:
                text = para.text
                colon_pos = text.find(":")
                
                if colon_pos > -1:  # If a colon is found
                    before_colon_text = text[:colon_pos].strip().upper()
                    after_colon_text = text[colon_pos + 1:].strip()
                    
                    # Replace HTML tags in the text after the colon with a single space
                    after_colon_text = re.sub(html_tag_pattern, " ", after_colon_text).strip()
                    
                    # Skip processing for certain keywords
                    if before_colon_text not in {"FILENAME", "BUTTON", "LABEL", "DIRECTIONS"}:
                        if after_colon_text:
                            # Add formatted content to the new document
                            if before_colon_text == "SCREEN TITLE":
                                para = new_doc.add_paragraph()
                                run = para.add_run(after_colon_text)
                                run.bold = True
                                run.underline = True
                            elif before_colon_text == "TITLE":
                                para = new_doc.add_paragraph()
                                run = para.add_run(after_colon_text)
                                run.bold = True
                            else:
                                new_doc.add_paragraph(after_colon_text)
                            content_added = True
            
            # Save the new document
            if content_added:
                new_doc.save(output_file)
                print(f"Content extracted and saved to '{output_file}'.")
            else:
                print("No relevant content found after colons in the document.")

        # Example usage
        input_docx = input_path
        output_docx = output_path
        extract_content_after_colon(input_docx, output_docx)

        # pass
        # doc = Document(input_path)
        # doc.add_paragraph("Processed with Content Builder Tool")
        # doc.save(output_path)

def html_tool(input_path, output_path):
    print(input_path , output_path)
    def extract_html_content(input_file, output_file):
        doc = Document(input_file)
        new_doc = Document()
        sections_content_pattern = re.compile(r"sections\.content", re.IGNORECASE)
        sections_header_title_pattern = re.compile(r"sections\.headerTitle", re.IGNORECASE)

        for table in doc.tables:
            for row in table.rows:
                left_cell_text = row.cells[0].text.strip()
                right_cell_text = row.cells[1].text.strip() if len(row.cells) > 1 else ""

                # Skip if the right cell is empty
                if not right_cell_text:
                    continue

                # Check if the left cell contains "sections.content..."
                if sections_content_pattern.search(left_cell_text):
                    para = new_doc.add_paragraph()
                    para.add_run(right_cell_text)

                # Check if the left cell contains "sections.headerTitle"
                if sections_header_title_pattern.search(left_cell_text):
                    para = new_doc.add_paragraph()
                    run = para.add_run(right_cell_text)
                    run.bold = True
                    run.underline = True
        
        new_doc.save(output_file)
        print(f'Sections content extracted and saved to {output_file}')

    # Example usage
    input_file = input_path
    output_file = output_path
    extract_html_content(input_file, output_file)


def validate_document(input_path, tool):
    """Check if the document contains required keywords for the tool."""
    doc = Document(input_path)
    required_keywords = KEYWORDS.get(tool, [])
    content = " ".join(paragraph.text for paragraph in doc.paragraphs)
    return any(keyword.lower() in content.lower() for keyword in required_keywords)

@app.route('/', methods=['GET'])
def index():
    return render_template('drag_drop.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    tool = request.form.get('tool')
    uploaded_file = request.files.get('file')
    force = request.form.get('force', 'false').lower() == 'true'
    # sections_to_remove = request.form.get('sections_to_remove', '').split(',')

    if not tool or not uploaded_file:
        return jsonify({"status": "error", "message": "Invalid request"}), 400

    # Save the uploaded file
    original_filename = uploaded_file.filename
    file_name, file_extension = os.path.splitext(original_filename)
    input_path = os.path.join(UPLOAD_FOLDER, uploaded_file.filename)
    uploaded_file.save(input_path)

    # Validate the document
    # if not validate_document(input_path, tool):
    #     return jsonify({"status": "error", "message": "The uploaded document is invalid for the selected option."})
    
        # Validate the document
    if not force and not validate_document(input_path, tool):
        return jsonify({
            "status": "error",
            "message": "The uploaded document is invalid for the selected option.",
            "force_option": True
        })


    # Process the document
    output_filename = f"{file_name}_after{file_extension}"
    output_path = os.path.join(OUTPUT_FOLDER, output_filename)

    if tool == 'rise_tool':
        rise_tool(input_path, output_path)
    elif tool == 'storyline_tool':
        storyline_tool(input_path, output_path)
    elif tool == 'content_builder_tool':
        content_builder_tool(input_path, output_path)
    elif tool == 'html_tool':
        html_tool(input_path, output_path)

    return jsonify({"status": "success", "message": "File processed successfully.", "download_url": f"/download/{output_filename}"})

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    file_path = os.path.join(OUTPUT_FOLDER, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return "File not found", 404

@app.route('/uploads/<filename>', methods=['POST'])
def cleanup_file(filename):
    try:
        processed_file_path = os.path.join(UPLOAD_FOLDER, filename)
        if os.path.exists(processed_file_path):
            os.remove(processed_file_path)
        return jsonify({"status": "success", "message": "File deleted successfully!"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

# Function to run Flask app
def run_flask():
    app.run(port=5000, debug=True, use_reloader=False)

# Expose Flask app using ngrok
public_url = ngrok.connect(5000).public_url
print(f"Public URL: {public_url}")

# Run Flask app in a separate thread
Thread(target=run_flask).start()
